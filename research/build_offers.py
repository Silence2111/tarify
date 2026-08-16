# -*- coding: utf-8 -*-
"""Гайд: подключение к CPA и выбор первых офферов (интернет + мобайл/SIM). .docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)
normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:cs'), 'Calibri')
ACCENT = RGBColor(0x1F, 0x4E, 0x79); GREEN = RGBColor(0x2E, 0x7D, 0x32); RED = RGBColor(0xC0, 0x39, 0x2B); GREY = RGBColor(0x60, 0x60, 0x60)


def set_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexc); tcPr.append(shd)


def para(text='', size=11, bold=False, italic=False, color=None, after=6, before=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
    return p


def rich(parts, size=11, after=6, before=0, color=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    for t, o in parts:
        r = p.add_run(t); r.font.size = Pt(o.get('size', size)); r.bold = o.get('bold', False); r.italic = o.get('italic', False)
        c = o.get('color', color)
        if c: r.font.color.rgb = c
    return p


def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(text).font.size = Pt(size); p.paragraph_format.space_after = Pt(3); return p


def num(text, size=11):
    p = doc.add_paragraph(style='List Number'); p.add_run(text).font.size = Pt(size); p.paragraph_format.space_after = Pt(3); return p


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs: r.font.color.rgb = ACCENT
    return p


def table(header, rows, fs=9, hfs=9):
    t = doc.add_table(rows=1, cols=len(header)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Light Grid Accent 1'
    hc = t.rows[0].cells
    for i, x in enumerate(header):
        hc[i].paragraphs[0].clear(); r = hc[i].paragraphs[0].add_run(x); r.bold = True; r.font.size = Pt(hfs)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); set_bg(hc[i], '1F4E79')
    for row in rows:
        cells = t.add_row().cells
        for i, x in enumerate(row):
            cells[i].paragraphs[0].clear(); r = cells[i].paragraphs[0].add_run(str(x)); r.font.size = Pt(fs)
    return t


# ===== ТИТУЛ =====
tt = doc.add_heading('Первые офферы', level=0)
for r in tt.runs: r.font.color.rgb = ACCENT
para('Подключение к CPA и выбор офферов: интернет + мобильные/SIM', size=14, italic=True, color=GREY, after=12)
rich([('Ставки — из проведённого исследования (карточки CPA-офферов Pampadu, категория A). '
       'Актуальные значения смотри в личном кабинете сети — они меняются.', {})], size=10, color=GREY, after=12)

# ===== ГЛАВНЫЙ ИНСАЙТ =====
h('0. Главный инсайт: начни с мобильных', level=1)
rich([('Мобильные тарифы / SIM-eSIM обходят самый сложный барьер — данные покрытия. ',
       {'bold': True, 'color': GREEN})])
para('Домашний интернет требует базы «какой провайдер в этом доме» — её собирать долго и дорого '
     '(главный ров, но и главный тормоз старта). А SIM/eSIM работает где угодно: адрес не нужен, '
     'можно промоутить выгодные тарифы по всей РФ сразу. eSIM вдобавок активируется онлайн за минуты '
     '— выше конверсия. Поэтому мобайл — самый быстрый путь к первым деньгам, а интернет собираешь '
     'параллельно.', after=4)
rich([('Минус мобайла: ', {'bold': True}), ('ниже чек (~280–550 ₽ против ~1300–2750 ₽ у интернета) — '
       'нужен объём. Плюс: нулевой барьер входа и национальный охват.', {})])

# ===== 1. CPA-СЕТИ =====
h('1. Какую CPA-сеть выбрать', level=1)
table(['Сеть', 'Плюсы', 'Минусы'],
      [['Pampadu', 'открытые витрины — ставки видны сразу, много телеком/мобайл офферов', 'меньше премиум-рекламодателей'],
       ['Admitad', 'крупнейшая, максимум офферов', 'телеком часто скрыт до модерации'],
       ['Cityads / Leads.su / Actionpay', 'дополнительные офферы', 'ставки видны после регистрации']])
rich([('Рекомендация: ', {'bold': True}),
      ('зарегистрируйся в Pampadu (ставки сразу видно) + Admitad (охват). Начни с Pampadu.', {})], before=4)

# ===== 2. РЕГИСТРАЦИЯ =====
h('2. Как зарегистрироваться (шаги)', level=1)
num('Зарегистрируйся как вебмастер/партнёр в сети.')
num('Добавь площадку — свой сайт tarify.rumedia-cdn.com (тип: сайт/веб).')
num('Пройди модерацию площадки (обычно быстро).')
num('Найди нужные офферы (интернет/мобайл), подай заявку на подключение к офферу.')
num('Получи партнёрскую ссылку / ID оффера — их будешь подставлять в заявки/кнопки на сайте.')
para('Заранее подготовь: описание площадки, тематику (сравнение тарифов связи), контакты. Это ускоряет '
     'модерацию.', size=10, italic=True, color=GREY, before=2)

# ===== 3. КРИТЕРИИ =====
h('3. Как выбирать первые офферы', level=1)
table(['Критерий', 'На что смотреть'],
      [['Ставка (₽)', 'сколько платят за подтверждённое действие'],
       ['Аппрув (AR)', 'доля подтверждённых; чем выше, тем лучше (телеком реально ~70% у Т-Мобайл)'],
       ['Холд', 'срок до подтверждения (~30 дней — норма)'],
       ['Целевое действие', 'что засчитывается: реальное подключение/платёж, а НЕ голая заявка'],
       ['Гео', 'доступность в твоём регионе (для интернета критично)'],
       ['eSIM/онлайн', 'для мобайла — предпочти офферы с онлайн-активацией (выше конверсия)']])

# ===== 4. ИНТЕРНЕТ =====
h('4. Домашний интернет — офферы', level=1)
table(['Провайдер', 'Ставка', 'Условия'],
      [['Ростелеком', 'от 2 752 ₽', 'оформление, холд 31 д'],
       ['Дом.ру', '~1 269–1 500 ₽', 'заявка/подключение'],
       ['МТС', '~1 800 ₽', 'подключение'],
       ['Уфанет (регион.)', 'до 1 000 ₽', 'за подключённого'],
       ['Диапазон рынка', '~1 300–2 750 ₽', 'за подтверждённое подключение']])
rich([('Барьер: ', {'bold': True, 'color': RED}),
      ('нужна база покрытия по адресам. Поэтому интернет — вторым шагом, по одному району за раз '
       '(см. data/COVERAGE.md).', {})], before=4)

# ===== 5. МОБАЙЛ =====
h('5. Мобильные тарифы / SIM — быстрый старт', level=1)
para('Ставки по SIM/eSIM (Pampadu, категория A):', bold=True, after=2)
table(['Оператор', 'Ставка', 'Детали'],
      [['Т-Мобайл (по этапам)', '553 + 510 + 1238 ₽', 'AR 71,7%; CR 1,16%; EPL 759 ₽; перенос номера +1238 ₽'],
       ['СберМобайл', '351 ₽ + RevShare 42,5%', 'оформление + пополнение от 150 ₽; % с пополнений 2-го мес.'],
       ['Билайн', '291 ₽ (фикс)', 'покупка SIM/eSIM'],
       ['Tele2 (t2)', 'до 280 ₽', 'покупка SIM'],
       ['ТВ-приставки Билайн', '9,7% от оплаты', 'сопутствующий оффер']])
para('Почему мобайл — хороший первый вертикал:', bold=True, before=4, after=2)
bullet('нет барьера покрытия — промоутишь по всей РФ, без адресной базы;', size=10)
bullet('eSIM активируется онлайн за минуты — короткая воронка, выше конверсия;', size=10)
bullet('RevShare (СберМобайл 42,5% с пополнений) — доход не разовый, а с оборота абонента;', size=10)
bullet('можно запустить страницу «выгодные тарифы SIM/eSIM» с партнёрскими ссылками почти без логики.', size=10)

# ===== 6. ТВ =====
h('6. ТВ и онлайн-кинотеатры — бонус-вертикаль', level=1)
table(['Сервис', 'Ставка'],
      [['Смотрёшка (онлайн-ТВ)', '11,3% от платежа + триал ~51 ₽'],
       ['Premier', '~461 ₽ за годовую подписку'],
       ['ivi / Кинопоиск / Okko / Start', '~50–460 ₽ за подписку/триал']])
para('Логично допродавать к интернету/мобайлу (bundle) или отдельным блоком «фильмы/ТВ».',
     size=10, italic=True, color=GREY, before=2)

# ===== 7. ПЛАТФОРМА =====
h('7. Как это ложится на платформу', level=1)
bullet('Модель уже поддерживает типы услуг: INTERNET / TV / MOBILE / BUNDLE (ServiceType). '
       'Мобильные тарифы можно завести тем же импортом, что и интернет.', size=10)
bullet('Мобайл НЕ требует привязки к дому — его можно показывать всем без адреса (отдельный раздел '
       '«мобильные тарифы», без матрицы покрытия).', size=10)
bullet('Заявка/партнёрская ссылка — тот же поток; для eSIM можно вести сразу на партнёрскую ссылку '
       'оффера (без обзвона).', size=10)
para('Небольшая доработка (когда дойдут руки): отдельная витрина мобильных тарифов + переход по '
     'партнёрской ссылке вместо заявки. Скажи — сделаю.', size=10, italic=True, color=GREY, before=2)

# ===== 8. ПЛАН ПЕРВОЙ НЕДЕЛИ =====
h('8. План первой недели', level=1)
num('Зарегистрируйся в Pampadu + Admitad, добавь площадку, пройди модерацию.')
num('Возьми 2–3 мобильных/eSIM оффера (Т-Мобайл, СберМобайл, Билайн) — они работают по всей РФ.')
num('Возьми 2–3 интернет-оффера для СВОЕГО города (Ростелеком/Дом.ру/МТС).')
num('Мобайл запусти сразу — нет барьера покрытия, собираешь первый трафик и конверсии.')
num('Параллельно собирай покрытие 1 района для интернета (data/COVERAGE.md), потом заливай CSV.')
num('Первые заявки по интернету обзванивай сам; по eSIM — веди на партнёрскую ссылку.')
para('Смысл: мобайл даёт быстрый старт и первые деньги без сбора покрытия; интернет — более высокий '
     'чек, но требует адресной базы, поэтому его подключаешь по району.', before=4, italic=True, color=GREY)

out = '/Users/Abdurahman/Projects/тарифы/Тарифы_первые_офферы.docx'
doc.save(out)
print('saved:', out)
