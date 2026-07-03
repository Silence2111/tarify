# -*- coding: utf-8 -*-
"""Сборка отчёта по исследованию тарифных платформ в .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- базовые стили ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn('w:cs'), 'Calibri')

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x60, 0x60, 0x60)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def para(text='', size=11, bold=False, italic=False, color=None, align=None,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def rich(parts, size=11, space_after=6, space_before=0, align=None, color=None):
    """parts: list of (text, {bold,italic,color}); color is a default for all parts"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for text, opt in parts:
        r = p.add_run(text)
        r.font.size = Pt(opt.get('size', size))
        r.bold = opt.get('bold', False)
        r.italic = opt.get('italic', False)
        c = opt.get('color', color)
        if c:
            r.font.color.rgb = c
    return p


def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


# ===================== ТИТУЛ =====================
title = doc.add_heading('Тарифные платформы', level=0)
for r in title.runs:
    r.font.color.rgb = ACCENT
para('Исследование рынка, коммерческая модель и рекомендации по MVP', size=14,
     italic=True, color=GREY, space_after=14)
rich([('Дата: ', {'bold': True}), ('15 июня 2026 г.', {}),
      ('    •    Стек проекта: ', {'bold': True}), ('Next.js + PostgreSQL', {})], size=10,
     color=GREY)
rich([('Методология: ', {'bold': True}),
      ('глубокое исследование с веб-источниками и адверсариальной проверкой фактов. '
       '5 направлений → 23 источника → 91 факт извлечён → 25 проверено '
       '(нужно 2 из 3 голосов «опровергнуть», чтобы факт убить) → 12 подтверждено, 13 убито. '
       '105 агентов, ~1.44 млн токенов.', {})], size=10, color=GREY, space_after=14)

doc.add_paragraph().add_run('').add_break()

# ===================== 1. КОММЕРЧЕСКАЯ МОДЕЛЬ =====================
h('1. Коммерческая модель — где деньги', level=1)
para('Пользователь ничего не платит. Платформа бесплатна для конечного пользователя, '
     'а платит провайдер — за приведённого клиента. Это lead-gen / CPA-модель.', space_after=8)

para('Два способа монетизации:', bold=True, space_after=4)
bullet('CPA (за подключение) — провайдер платит фикс за каждого реально подключённого абонента. '
       'Основной и самый «жирный» вариант для домашнего интернета и ТВ.')
bullet('CPL (за лид) — платят за квалифицированную заявку (оставил телефон, согласился на звонок), '
       'даже если не подключился. Дешевле за штуку, но объём больше.')

para('Потолок модели (подтверждено исследованием):', bold=True, space_before=6, space_after=4)
para('MoneySuperMarket на этой же модели делает £446 млн выручки и ~18% чистой маржи. '
     'Маржа высокая именно потому, что нет товара и склада — только трафик и передача заявок; '
     'себестоимость единицы близка к нулю (один сайт обслуживает и 100, и 100 000 заявок).')

rich([('Честная дыра в данных: ', {'bold': True, 'color': RED}),
      ('сколько именно платит российский провайдер за подключённого абонента — исследование НЕ '
       'подтвердило. Вся проверенная юнит-экономика британская (~£22–30 за переключение энергии, '
       '2015 г.). Это вынесено в «открытые вопросы» как главный неизвестный параметр.', {})],
     space_before=6)

para('Ориентир из общих знаний рынка РФ (НЕ из проверенного отчёта, требует валидации):', bold=True,
     space_before=6, space_after=4)
para('В CPA-сетях типа Admitad подключение домашнего интернета обычно стоит провайдеру '
     'ориентировочно ~500–2000 ₽ за абонента; мобильные SIM — заметно меньше.')

# формула
fp = para('', space_before=4, space_after=4)
fr = fp.add_run('Выручка = трафик × конверсия в заявку × конверсия заявки в подключение × CPA\n'
                'Пример (иллюстрация, не факт):\n'
                '10 000 визитов/мес × 4% в заявку × 25% в подключение × 1000 ₽ = 100 000 ₽/мес')
fr.font.name = 'Consolas'
fr.font.size = Pt(10)
fr.font.color.rgb = ACCENT

para('Узкие места — трафик (SEO/контекст) и процент подтверждённых подключений (за это и платит '
     'провайдер). Поэтому первый MVP нужен, чтобы померить эти две конверсии на реальных данных, '
     'а не гадать.', space_before=4)
para('Дополнительные источники денег той же платформы: реклама провайдеров (баннер/приоритет в '
     'выдаче), позже — продажа B2B-SaaS «конструктор тарифов» провайдерам (направление 3/4).')

# ===================== 2. СВОДКА =====================
h('2. Сводка (вердикт исследования)', level=1)
para('Из четырёх направлений проверенные данные указывают: сравнение телеком/интернета и '
     'SaaS-биллинг / usage-based pricing — наиболее валидированные бизнес-модели; ЖКХ/энергетика '
     'упирается в жёсткую регуляторную стену в РФ.')
para('Агрегаторы (Tarifnik.ru в РФ; broadbandchoices, MoneySuperMarket, uSwitch в UK) '
     'монетизируются через комиссии за лид/переключение — модель доказана и масштабируется до '
     '~£446 млн выручки и ~18% чистой маржи (MoneySuperMarket), но комиссии за лид скромные '
     '(~£22–30) и объёмы переключений волатильны (в 2021 в UK переключения рухнули, 25+ '
     'энергокомпаний обанкротились).')
para('Розничный рынок электроэнергии РФ — гос-регулируемые тарифы и «гарантирующие поставщики» без '
     'механизма смены для домохозяйств → uSwitch-модель для физлиц в РФ структурно невозможна '
     '(B2B частично либерализован).')
para('SaaS-биллинг — сильный, но очень тесный и капиталоёмкий глобальный рынок (Stripe Billing, '
     'Zuora, Chargebee, Orb, Metronome — последнюю Stripe купил за $1 млрд в дек-2025). Для '
     'одиночки/малой команды в СНГ самый валидированный и не занятый коридор — РФ/СНГ агрегатор '
     'телеком+интернет+ТВ; «универсальный конструктор биллинга» перспективен, но капитало- и '
     'конкуренто-ёмок, и входить в него лучше через узкую вертикаль или нишу B2B-закупки энергии.')

# ===================== 3. ТАБЛИЦА =====================
h('3. Сравнительная таблица направлений', level=1)
rows = [
    ('Направление', 'Потенциал', 'Сложность входа', 'Монетизация', 'Ниша в СНГ', 'Вердикт'),
    ('1. Связь/интернет/ТВ', 'Средний–высокий', 'Низкая', 'Комиссия за лид/подключение',
     'Не насыщена', '🟢 Начать здесь'),
    ('2. ЖКХ/энергетика (физлица РФ)', 'Низкий', 'Высокая (регуляторика)', '—',
     'Заблокирована', '🔴 Пропустить'),
    ('2b. B2B-закупка энергии', 'Средний', 'Средняя', 'Комиссия/брокеридж',
     'Почти свободна', '🟡 Открыта'),
    ('3/4. SaaS-биллинг (горизонталь)', 'Высокий', 'Очень высокая', 'SaaS + % с транзакций',
     'Глобально тесно', '🟡 Через нишу'),
]
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, txt in enumerate(rows[0]):
    hdr[i].paragraphs[0].clear()
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(hdr[i], '1F4E79')
for row in rows[1:]:
    cells = table.add_row().cells
    for i, txt in enumerate(row):
        cells[i].paragraphs[0].clear()
        r = cells[i].paragraphs[0].add_run(txt)
        r.font.size = Pt(9)

# ===================== 4. ПОДТВЕРЖДЁННЫЕ НАХОДКИ =====================
h('4. Подтверждённые находки (4)', level=1)

findings = [
    {
        'title': '1. Агрегаторы сравнения телеком/интернет/ТВ — доказанный lead-gen бизнес в РФ и на '
                 'Западе; монетизация через комиссии провайдеров, не подписку юзера; ниша в СНГ '
                 'активна, но не насыщена.',
        'meta': 'confidence: high · голоса 3-0 (Tarifnik), 3-0 (broadbandchoices)',
        'ev': 'Tarifnik.ru — бесплатный для пользователя агрегатор домашнего интернета, ТВ и '
              'мобильной связи в РФ, прямо заявляет выручку от «партнёрских договоров с провайдерами '
              'услуг связи» как «официальный партнёр более чем 100 провайдеров», с колл-центром и '
              'обратным звонком (оператор: ИП Горбушин А.А.); на 2GIS независимо помечен как '
              '«агрегатор интернет-провайдеров». broadbandchoices.co.uk — идентичная модель в UK '
              '(broadband, телефон, ТВ, мобайл), выручка «advertising fees and commissions from a '
              'number of the providers», аффилиат «до £30 за продажу». Есть рабочий РФ-образец, но '
              'поле не переполнено — самый доступный валидированный коридор для малой команды СНГ.',
        'src': ['tarifnik.ru', 'startupintros.com/orgs/broadbandchoices-co-uk',
                'broadbandchoices.co.uk/about-us', 'broadbandchoices.co.uk/solutions-for-affiliates'],
    },
    {
        'title': '2. Модель сравнения/lead-gen масштабируется до большой выручки и высокой маржи в '
                 'зрелости, но комиссии за лид скромные, а объёмы переключений волатильны/цикличны.',
        'meta': 'confidence: high · голоса 2-1 (выручка MSM), 3-0 (модель MSM), 2-1 (£30/£60), '
                '3-0 (£29/£30 за переключение), 3-0 (банкротства)',
        'ev': 'MoneySuperMarket (MONY Group plc) за FY2025: выручка £446.3 млн, операционная прибыль '
              '£117.4 млн, чистая прибыль ~£80.7 млн (~18% маржа) — подтверждено первичной '
              'RNS-отчётностью; это потолок зрелого агрегатора. Ядро модели — комиссия/реферальные '
              'сборы с провайдеров за покупку клиента. Юнит-экономика лида скромная: энергетические '
              'PCW в UK зарабатывали ~£22–30 за переключение одного топлива (MoneySuperMarket ~£29, '
              'uSwitch ~£30), ~£44–60 за dual-fuel. Модель цикличная и хрупкая в энергетике: 25+ '
              'энергокомпаний UK обанкротились в 2021 на скачке оптовых цен (NAO: 29 поставщиков за '
              'год, ~4 млн домохозяйств, £2.7 млрд). Оговорка: £30/£60 — бенчмарк 2015 г., сейчас '
              'вероятно ниже.',
        'src': ['en.wikipedia.org/wiki/MoneySuperMarket',
                'financialadvice.co.uk/...price-comparison-websites-charge-30-per-switch',
                'power.nridigital.com/...uk_broken_energy_market',
                'nao.org.uk/reports/the-energy-supplier-market'],
    },
    {
        'title': '3. uSwitch-подобная платформа сравнения и переключения энергии для физлиц '
                 'структурно заблокирована в РФ: гос-регулируемые тарифы, гарантирующие поставщики, '
                 'отсутствие механизма смены.',
        'meta': 'confidence: high · голоса 2-1, 2-1',
        'ev': 'Фед. закон РФ: «Гарантирующие поставщики будут осуществлять поставку электроэнергии '
              'на основе регулируемых тарифов», розничная цена = опт + передача/распределение + '
              'регулируемая сбытовая надбавка — гос-композит, не конкурентная цена. Домохозяйства '
              'обслуживаются строго по регулируемым тарифам региональных властей, бытовой рынок '
              '«всё ещё не открыт» (рецензируемый ScienceDirect), есть «отсутствие механизма смены '
              'гарантирующего поставщика». Это убирает ценовую конкуренцию и переключения, на '
              'которых держатся западные энергетические PCW. Лазейка: небытовые/B2B потребители '
              'платят свободную цену и МОГУТ менять энергосбыт с ~2011 г. → ниша B2B-сравнения '
              'закупки энергии существует. Направление 2 — самое слабое потребительское направление '
              'в РФ.',
        'src': ['consultant.ru/document/cons_doc_LAW_32472', 'np-sr.ru', 'mosenergosbyt.ru'],
    },
    {
        'title': '4. Направление SaaS-биллинг / usage-based — реальное и растущее (рынок решительно '
                 'уходит от flat-rate к гибриду подписка+потребление), но это тесный и сильно '
                 'капитализированный глобальный ландшафт.',
        'meta': 'confidence: high · голоса 3-0 (сдвиг к гибриду), 3-0 (тесный ландшафт), '
                '3-0 (покупка Metronome)',
        'ev': 'Сдвиг ценообразования — консенсус индустрии: OpenView — 61% SaaS-компаний на '
              'гибридном pricing (было 49% в 2024), usage-based ~27% (2021) → ~38% (2026), '
              'доминирующий паттерн «базовая подписка + overage по потреблению»; подтверждают Maxio, '
              'Flexera, Chargebee. Поле (направления 3 и 4) тесное: Stripe Billing (~32–34% доли '
              'payment-management, оценка ~$70–91 млрд), Zuora (куплена за $1.7 млрд, окт-2024), '
              'Chargebee ($479.9 млн привлечено, 6500+ клиентов), плюс Orb, Amberflo, M3ter и '
              'Metronome — которую Stripe купил за $1 млрд в дек-2025 (2.1x к приватной оценке '
              '$470 млн). Спрос на инфраструктуру биллинга подтверждён, но сигнализирует о высокой '
              'конкуренции и капиталоёмкости для горизонтального «universal billing-as-a-service».',
        'src': ['globenewswire.com/...Subscription-Billing-Management-Market-Report-2026',
                'research.contrary.com/company/metronome', 'sacra.com/c/metronome',
                'stripe.com/newsroom/news/stripe-completes-metronome-acquisition'],
    },
]

for f in findings:
    rich([('✓ ', {'bold': True, 'color': GREEN}), (f['title'], {'bold': True})], space_before=8,
         space_after=2)
    para(f['meta'], size=9, italic=True, color=GREY, space_after=4)
    para(f['ev'], size=10, space_after=4)
    rich([('Источники: ', {'bold': True, 'size': 9}),
          (' · '.join(f['src']), {'size': 9, 'color': GREY})], space_after=4)

# ===================== 5. УБИТЫЕ УТВЕРЖДЕНИЯ =====================
h('5. Убитые проверкой утверждения (13)', level=1)
para('Эти заявления не прошли адверсариальную проверку (голоса = за-опровержение). Именно поэтому в '
     'рекомендациях нет конкретных TAM-цифр — почти все «рынок $N млрд» взяты из пресс-релизов и '
     'блогов и проверку не прошли.', space_after=6)
killed = [
    ('broadbandchoices ~$10 млн годовой выручки в 2024', '0-3'),
    ('«Сложный биллинг — топ-барьер подписок, 26% SMB» (как обоснование рынка BaaS)', '0-3'),
    ('Обвал переключений в UK в дек-2021 до ⅓ от 2020', '1-2'),
    ('ЖКХ-тарифы в РФ задаются на муниципальном уровне, нет единого нац. источника', '0-3'),
    ('«Нет универсального ЖКХ-калькулятора, данные не агрегированы»', '1-2'),
    ('Usage-based pricing у 85% крупных софт-компаний (с 50% в 2021)', '0-3'),
    ('Рынок subscription billing $8.47 млрд (2025) → $37.36 млрд (2035), CAGR 16%', '1-2'),
    ('Рынок billing/invoicing софта $4.3 млрд (2024) → $13.1 млрд (2031), 15.2%', '0-3'),
    ('Usage-based рос с 27% (2018) до 61% (2022)', '0-3'),
    ('Convergent billing $20.98 млрд (2026) → $33.24 млрд (2031), 9.66%', '0-3'),
    ('5 вендоров (Amdocs/Oracle/Ericsson/Huawei/CSG) держат ~58% convergent billing', '0-3'),
    ('Recurly берёт 0.9% со всей выручки + месячная плата', '0-3'),
    ('Chargebee: месячная плата + % с выручки сверх лимита, free-план Launch', '0-3'),
]
for txt, vote in killed:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(txt + '  ')
    r1.font.size = Pt(10)
    r2 = p.add_run('[' + vote + ']')
    r2.font.size = Pt(10)
    r2.bold = True
    r2.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(2)

# ===================== 6. ОТКРЫТЫЕ ВОПРОСЫ =====================
h('6. Открытые вопросы (доисследовать перед стартом)', level=1)
oq = [
    'Реальная юнит-экономика РФ/СНГ-агрегатора (CAC, комиссия за лид, конверсии) — сколько РФ-'
    'провайдер платит за подключённого абонента (вся проверенная комиссия — британская).',
    'Насколько велик и доступен частично-либерализованный сегмент B2B-закупки энергии в РФ '
    '(единственный открытый коридор энергетики).',
    'Защитимая bottom-up оценка размера рынка для (а) СНГ-ниши сравнения телеком/интернета и '
    '(б) вертикального конструктора тарифов для МСБ СНГ — без пресс-релизных TAM.',
    'Есть ли незанятая ниша «универсального конструктора тарифов» именно под регулируемые/utility-'
    'тарифы (ЖКХ, телеком, транспорт) в СНГ — B2B-SaaS для настройки и расчёта сложных ступенчатых '
    'тарифов, а не лобовая конкуренция со Stripe/Zuora/Chargebee.',
]
for i, q in enumerate(oq, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(q).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

# ===================== 7. РЕКОМЕНДАЦИИ И MVP =====================
h('7. Рекомендации и MVP', level=1)
para('Куда двигаться одиночке / маленькой команде:', bold=True, space_after=4)
para('Стартовать с направления 1 — агрегатор сравнения связи/интернета/ТВ для СНГ. Модель доказана '
     '(Tarifnik), поле не насыщено, монетизация понятна (комиссия с провайдеров за подключение), '
     'капитал минимальный, регуляторных стен нет.')

para('Менее очевидные, но перспективные ниши:', bold=True, space_before=6, space_after=4)
bullet('Вертикальный конструктор тарифов для регулируемых отраслей (ЖКХ, телеком, транспорт, '
       'парковки) — B2B-SaaS, помогающий поставщикам настраивать и считать сложные ступенчатые/'
       'зонные тарифы. Не лобовая конкуренция со Stripe/Zuora.')
bullet('B2B-сравнение закупки энергии — единственная открытая энергетическая лазейка в РФ.')
bullet('Tariff-as-content / SEO-актив — публичный калькулятор/сравнение тарифов генерит органический '
       'трафик, монетизируемый лидами (так растут Tarifnik / MoneySuperMarket).')

para('MVP, который имеет смысл собрать первым:', bold=True, space_before=6, space_after=4)
para('Агрегатор тарифов связи/интернета для одного города/региона:')
bullet('модель данных «тариф» (провайдер → план → опции → цены → правила) — общее ядро для всех '
       'направлений, не выбросим при пивоте;')
bullet('сравнение/фильтрация по адресу и параметрам;')
bullet('захват лида (форма «перезвоните мне») → передача провайдеру;')
bullet('старт с ручного парсинга 5–10 провайдеров, без интеграций.')
para('Это проверяет и спрос пользователей, и готовность провайдеров платить комиссию — за пару '
     'недель на Next.js + PostgreSQL.', space_before=2)

# ===================== 8. ИСТОЧНИКИ =====================
h('8. Источники (23, с оценкой качества)', level=1)
sources = {
    'Связь/интернет': ['tarifnik.ru (primary)', 'financialadvice.co.uk (secondary)',
                       'MoneySuperMarket/Wikipedia (secondary)', 'startupintros (secondary)',
                       'businessresearchinsights (unreliable)', 'tadviser/Сравни.ру (unreliable)'],
    'Энергетика/ЖКХ': ['consultant.ru/закон (primary)', 'gogov.ru (secondary)',
                       'thefintechtimes/uSwitch (secondary)', 'nridigital (secondary)',
                       'cornwall-insight (unreliable)', 'dealroom/uswitch (unreliable)'],
    'SaaS-биллинг': ['sacra/metronome (secondary)', 'globenewswire (secondary)',
                     'contrary/metronome (secondary)', 'zeni.ai (blog)'],
    'Billing-as-a-service': ['mordorintelligence (secondary)', 'withorb/chargebee (blog)',
                             'dodopayments (blog)', 'grandviewresearch (unreliable)'],
    'Стратегия/ниши': ['entrepreneurloop (blog)', 'f3fundit (blog)',
                       'vc.ru/свободные ниши (blog)'],
}
for cat, items in sources.items():
    para(cat, bold=True, size=10, space_before=4, space_after=2)
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(it).font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)

# ===================== 9. СТАТИСТИКА =====================
h('9. Статистика прогона', level=1)
stats = [
    ('Углов исследования', '5'),
    ('Источников загружено', '23'),
    ('Фактов извлечено', '91'),
    ('Фактов проверено', '25'),
    ('Подтверждено', '12'),
    ('Убито проверкой', '13'),
    ('После синтеза', '4'),
    ('Отброшено по бюджету', '6'),
    ('Агентов', '105'),
    ('Токенов', '~1.44 млн'),
    ('Время работы', '~30 мин'),
]
st = doc.add_table(rows=0, cols=2)
st.style = 'Light List Accent 1'
for k, v in stats:
    c = st.add_row().cells
    rk = c[0].paragraphs[0].add_run(k)
    rk.font.size = Pt(10)
    rk.bold = True
    c[1].paragraphs[0].add_run(v).font.size = Pt(10)

# ===================== оговорки =====================
h('Оговорки', level=2)
para('Юнит-экономика — на данных UK, для СНГ напрямую не проверена (главное неизвестное — ставка CPA '
     'российских провайдеров). Цифры TAM из пресс-релизов отсеяны как недостоверные. Энергокомиссии '
     '£22–30 — бенчмарк 2015 г.; объёмы переключений в UK рухнули после 2021. Оценка Stripe ($70 млрд) '
     'устарела (позже ~$91.5 млрд), но вывода о тесноте рынка не меняет.', size=10, italic=True,
     color=GREY)

out = '/Users/Abdurahman/Projects/тарифы/Тарифные_платформы_исследование.docx'
doc.save(out)
print('saved:', out)
